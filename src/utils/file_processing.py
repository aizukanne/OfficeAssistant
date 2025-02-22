import base64
import io
import json
import os
import tempfile
from typing import Optional, Dict, Any, BinaryIO, Union
import wave

import PyPDF2
import docx
import markdown2
import requests
from bs4 import BeautifulSoup

def read_pdf(file_data: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts text from a PDF file.
    
    Args:
        file_data: PDF file data as string path, bytes, or file object
        
    Returns:
        str: Extracted text
    """
    try:
        # Create PDF reader object
        if isinstance(file_data, str):
            # It's a file path
            pdf = PyPDF2.PdfReader(open(file_data, 'rb'))
        elif isinstance(file_data, (bytes, io.BytesIO)):
            # It's bytes or BytesIO
            pdf = PyPDF2.PdfReader(io.BytesIO(file_data) if isinstance(file_data, bytes) else file_data)
        else:
            # It's a file object
            pdf = PyPDF2.PdfReader(file_data)

        # Extract text from all pages
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
            
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def read_docx(file_data: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts text from a DOCX file.
    
    Args:
        file_data: DOCX file data as string path, bytes, or file object
        
    Returns:
        str: Extracted text
    """
    try:
        # Create Document object
        if isinstance(file_data, str):
            # It's a file path
            doc = docx.Document(file_data)
        else:
            # It's bytes or file object
            doc = docx.Document(io.BytesIO(file_data) if isinstance(file_data, bytes) else file_data)

        # Extract text from paragraphs
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def read_markdown(file_data: Union[str, bytes]) -> str:
    """
    Converts Markdown to plain text.
    
    Args:
        file_data: Markdown content as string or bytes
        
    Returns:
        str: Plain text
    """
    try:
        # Convert to string if needed
        if isinstance(file_data, bytes):
            file_data = file_data.decode('utf-8')
            
        # Convert Markdown to HTML
        html = markdown2.markdown(file_data)
        
        # Convert HTML to plain text
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    except Exception as e:
        print(f"Error reading Markdown: {e}")
        return ""

def download_and_read_file(file_url: str, mime_type: str) -> str:
    """
    Downloads and reads content from a file URL.
    
    Args:
        file_url: URL of the file
        mime_type: MIME type of the file
        
    Returns:
        str: File content
    """
    try:
        # Download file
        response = requests.get(file_url)
        response.raise_for_status()
        file_data = response.content
        
        # Process based on MIME type
        if mime_type.startswith('application/pdf'):
            return read_pdf(file_data)
        elif mime_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            return read_docx(file_data)
        elif mime_type.startswith('text/markdown'):
            return read_markdown(file_data)
        elif mime_type.startswith('text/'):
            return file_data.decode('utf-8')
        else:
            return f"Unsupported file type: {mime_type}"
    except Exception as e:
        print(f"Error downloading/reading file: {e}")
        return ""

def convert_to_wav_in_memory(audio_data: bytes) -> bytes:
    """
    Converts audio data to WAV format in memory.
    
    Args:
        audio_data: Audio file data
        
    Returns:
        bytes: WAV format audio data
    """
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.m4a', delete=False) as temp_in, \
             tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_out:
            
            # Write input data
            temp_in.write(audio_data)
            temp_in.flush()
            
            # Convert using ffmpeg
            os.system(f'ffmpeg -i {temp_in.name} {temp_out.name}')
            
            # Read output data
            with open(temp_out.name, 'rb') as f:
                wav_data = f.read()
                
        # Clean up temporary files
        os.unlink(temp_in.name)
        os.unlink(temp_out.name)
        
        return wav_data
    except Exception as e:
        print(f"Error converting audio: {e}")
        return b""

def create_pdf(text: str) -> bytes:
    """
    Creates a PDF from text.
    
    Args:
        text: Text to convert to PDF
        
    Returns:
        bytes: PDF file data
    """
    try:
        from fpdf import FPDF
        
        # Create PDF object
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add text
        pdf.multi_cell(0, 10, text)
        
        # Get PDF data
        return pdf.output(dest='S').encode('latin1')
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return b""

def get_file_extension(filename: str) -> str:
    """
    Gets the extension from a filename.
    
    Args:
        filename: Name of the file
        
    Returns:
        str: File extension
    """
    return os.path.splitext(filename)[1].lower()

def get_mime_type(filename: str) -> str:
    """
    Gets the MIME type for a file.
    
    Args:
        filename: Name of the file
        
    Returns:
        str: MIME type
    """
    import mimetypes
    return mimetypes.guess_type(filename)[0] or 'application/octet-stream'

def is_valid_file_type(filename: str, allowed_types: list) -> bool:
    """
    Checks if a file type is allowed.
    
    Args:
        filename: Name of the file
        allowed_types: List of allowed extensions
        
    Returns:
        bool: Whether the file type is allowed
    """
    ext = get_file_extension(filename)
    return ext.lower() in [t.lower() for t in allowed_types]

def get_file_size(file_data: Union[str, bytes, BinaryIO]) -> int:
    """
    Gets the size of a file.
    
    Args:
        file_data: File data as string path, bytes, or file object
        
    Returns:
        int: File size in bytes
    """
    if isinstance(file_data, str):
        return os.path.getsize(file_data)
    elif isinstance(file_data, bytes):
        return len(file_data)
    else:
        pos = file_data.tell()
        size = file_data.seek(0, 2)
        file_data.seek(pos)
        return size